from datetime import datetime
import os
import shutil
from flask import Flask,render_template, request, redirect, url_for, session, flash, jsonify
import fitz
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import CountVectorizer
import hashlib
import numpy as np
from docx import Document
from flask_mysqldb import MySQL
from uuid import uuid4
import MySQLdb.cursors
import re
import cv2
import io
import nltk
import imghdr
import pdfplumber

nltk.download('wordnet')

app = Flask(__name__, static_folder="D:\Kuliah\jajal\Development Manpro_Rev\static")

app.secret_key="hehehayoapa"

app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = 'proyekmanpro'
app.config['MYSQL_DB'] = 'dbmanpro'
 
mysql = MySQL(app)

# 
allowedExt={'docx', 'pdf', 'jpg'}
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowedExt
def is_user_logged_in():
    return 'user_id' in session

# @app.route('/uploaded')
def is_file_uploaded():
    return 'doc_id' in session

@app.route('/hello')
def hello():
    return 'Hello World'

@app.route('/home')
def home():
    if is_user_logged_in():
        # User is logged in, render template with logout button
        return render_template('index.html', logged_in=True)
    else:
        # User is not logged in, render template with login button
        return render_template('index.html', logged_in=False)


@app.route('/uploadform')
def uploadform():
    if is_user_logged_in() and is_file_uploaded():
        # User is not logged in, render template with login button
        return render_template('upload.html', logged_in=True, uploaded = True)
    elif is_user_logged_in() and not is_file_uploaded():
        # User is logged in, render template with logout button
        return render_template('upload.html', logged_in=True, uploaded = False)
    else:
        # User is not logged in, render template with login button
        return render_template('upload.html', logged_in=False, uploaded = False)


 
@app.route('/signup', methods = ['GET','POST'])
def signup():
    if request.method == 'POST' and 'signup-username' in request.form and 'signup-email' in request.form and 'signup-password' in request.form:
        user_id = str(uuid4())
        username = request.form['signup-username']
        email = request.form['signup-email']
        password = request.form['signup-password']
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute("SELECT * FROM user WHERE user_email = %s", (email,))
        account= cursor.fetchone()
        if account:
            acc_err="Account already exists"
            return jsonify(error=acc_err)
        elif not re.match(r'[^@]+@[^@]+\.[^@]+',email):
            email_err="Invalid email address"
            return jsonify(error=email_err)
        elif not re.match(r'^[a-zA-Z0-9]+$',username):
            username_err="Username must contain only characters and numbers"
            return jsonify(error=username_err)
        elif not username or not password or not email:
            form_err="Please fill out the form"
            return jsonify(error=form_err)
        else:
            cursor.execute("""INSERT INTO user VALUES(%s,%s,%s,%s)""",(user_id,username,email,password,))
            mysql.connection.commit()
            msg="You have successfully signed up"
            cursor.close()
            return jsonify(success = True)
        
    return redirect('/home')

@app.route('/login', methods = ['GET','POST'])
def login():
    msg = ''
    error= ''
    if request.method == 'POST' and 'signin-email' in request.form and 'signin-password' in request.form:
        email = request.form['signin-email']
        password = request.form['signin-password']
        cursor = mysql.connection.cursor(MySQLdb.cursors.DictCursor)
        cursor.execute("SELECT * FROM user WHERE user_email = %s and user_pass = %s", (email, password,))
        account = cursor.fetchone()
        if account:
            session['loggedin'] = True
            session['user_id'] = account['user_id']
            session['user_name'] = account['user_name']
            msg="You have successfully logged in"
            return jsonify(success = True, logged_in = True)
            # return render_template('index.html', msg=msg, logged_in = True)
        else:
            error = "Incorrect password/username"
            return jsonify(error = error, logged_in = False)
            # return render_template('index.html', error = error, logged_in = False)
    # return render_template('index.html',session=session, msg=msg, logged_in = True)
@app.route('/test')
def test():
    return session.get('user_id', 'no user id')
@app.route('/logout')
def logout():
    session.pop('loggedin', None)
    session.pop('user_id', None)
    session.pop('user_name', None)
    return(redirect('/home'))

@app.route('/upload', methods=['GET', 'POST'])
def upload_file():
    file = request.files['file']
    filename = file.filename
    docid = str(uuid4())
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    if file and allowed_file(filename):
        ext = filename.rsplit('.', 1)[1].lower()
        
        #APABILA EXT PDF MAKA PERLU DI EKSTRAK TEXT2DOCX DAN IMAGE2PNG
        if ext == 'pdf':
            temp_folder = f"temp/{docid}"  # Folder temporary untuk menyimpan file PDF sementara
            os.makedirs(temp_folder)
            jpg_path = f"{temp_folder}/{filename}"
            file.save(jpg_path)
            text = extract_text_from_pdf(jpg_path)
            
            docContentNew = create_docx_from_text(text)
            # Melakukan pengolahan teks dan gambar PDF sesuai kebutuhan Anda
            # ...
            # Mengganti nama file sebelum ekstensi menjadi 'extPDF.docx'
            new_filename = os.path.splitext(filename)[0] + '_PDF.docx'
            if 'user_id' in session:
                cursor = mysql.connection.cursor()
                cursor.execute("INSERT INTO document (document_id, user_id, doc_name, doc_content, timestamp) VALUES (%s, %s, %s, %s, %s)",
               (docid, session['user_id'], new_filename, docContentNew, current_time,))
                mysql.connection.commit()
                # Setelah selesai pengolahan, hapus folder temporary
                extract_images_from_pdf(jpg_path, temp_folder, docid, new_filename)
                shutil.rmtree(temp_folder)
                if 'doc_id' in session:
                    session.pop('doc_id', None)
                    session['doc_id'] = docid
                else:
                    session['doc_id'] = docid
                return jsonify(success=True, uploaded=True)
            else:
                # Setelah selesai pengolahan, hapus folder temporary
                shutil.rmtree(temp_folder)
                msg = 'Please login to upload file'
                return jsonify(error=msg, uploaded=False)
        elif ext == 'jpg':
            if 'user_id' in session:
                cursor = mysql.connection.cursor()
                cursor.execute("INSERT INTO image_doc (img_id, image_name, image_file) VALUES (%s, %s, %s)",
                        (docid, filename, file,))
                mysql.connection.commit()
                if 'doc_id' in session:
                    session.pop('doc_id', None)
                    session['doc_id'] = docid
                else:
                    session['doc-id'] = docid
                return jsonify(success=True, uploaded=True)
            
            
            else:
                # Setelah selesai pengolahan, hapus folder temporary
                shutil.rmtree(temp_folder)
                msg = 'Please login to upload file'
                return jsonify(error=msg, uploaded=False)
            
            
        else:
            docContent = file.read()
            if 'user_id' in session:
                cursor = mysql.connection.cursor()
                cursor.execute("INSERT INTO document (document_id, user_id, doc_name, doc_content, timestamp) VALUES (%s, %s, %s, %s, %s)",
               (docid, session['user_id'], file.filename, docContent, current_time,))
                if 'doc_id' in session:
                    session.pop('doc_id', None)
                    session['doc_id'] = docid
                else:
                    session['doc_id'] = docid
                return jsonify(success=True, uploaded=True)
            else:
                msg = 'Please login to upload file'
                return jsonify(error=msg, uploaded=False)
    else:
        msg = 'Invalid file type'
        return jsonify(error=msg, uploaded=False)
    
def create_docx_from_text(text):
    doc = Document()
    doc.add_paragraph(text)
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream.getvalue()

def extract_text_from_pdf(jpg_path):
    doc = fitz.open(jpg_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text

def extract_images_from_pdf(jpg_path, output_folder, doc_id, new_filename):
    doc = fitz.open(jpg_path)
    for i, page in enumerate(doc):
        image_list = page.get_images()
        for image_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_name = f"halaman{i}_gambar_{image_index}_{new_filename}.png"
            image_path = f"{output_folder}/{image_name}"
            with open(image_path, "wb") as image_file:
                image_file.write(base_image["image"])
            
            # Insert data gambar ke tabel image_doc
            cursor = mysql.connection.cursor()
            cursor.execute("INSERT INTO image_doc (img_id, doc_id, image_name, image_file) VALUES (%s, %s, %s, %s)",
                           (str(uuid4()), doc_id, image_name, base_image["image"],))
            mysql.connection.commit()
    
def jaccard_similarity(set1, set2):
    intersection = len(set1.intersection(set2))
    union = len(set1.union(set2))
    return intersection / union

def average_similarity(set1, set2):
    # set1 = set(list1)
    # set2 = set(list2)
    intersection = len(set1.intersection(set2))
    # union = set1.union(set2)
    return intersection / len(set2)

def preprocess_sentences(text):
    text = re.sub(r'["“”]', '', text)
    text = text.replace('\n', ' ')
    text = ' '.join(text.split())
    text = re.sub(' +', ' ', text)
    text = re.sub(r'[^\w\s]+$', '', text)
    text = re.sub(r'\s*([^\w\s])\s*', r' \1 ', text)
    text = re.sub(r'\s([^\w\s])\s', r'\1 ', text)
    sentences = [sentence.strip() for sentence in text.split('.')]
    return sentences

@app.route('/jaccard_similarity', methods=['GET','POST'])
def calculate_jaccard_similarity():
    # file1 = request.files['tes.docx']
    # file2 = request.files['tes2.docx']
    # doc_id1 = '9c798b43-1d9e-40db-9e7b-2d6144d81ded'
    # doc_id2 = 'd1512997-78ad-4b91-a296-3afc57b35c19'
    doc_id = session['doc_id']

    cursor = mysql.connection.cursor()
    cursor.execute("SELECT doc_content FROM document WHERE document_id = %s", (doc_id,))
    row = cursor.fetchone()

    if row:
        content1 = row[0]
        if imghdr.what(None, h=content1) is not None:
            return 'Invalid document type'
        if content1.startswith(b'%PDF'):
            pdf = pdfplumber.open(io.BytesIO(content1))
            extracted_text = ""
            for page in pdf.pages:
                extracted_text += page.extract_text()
            paragraf_docx = extracted_text.lower()
        else:
            paragraf_docx = ""
            document = Document(io.BytesIO(content1))
            for paragraph in document.paragraphs:
                paragraf_docx += paragraph.text.lower()
            print(paragraf_docx)

        cursor.execute("SELECT document_id, doc_content FROM document WHERE document_id != %s", (doc_id,))
        similarity_results_jaccard = []
        similarity_results_average = []
        for row in cursor.fetchall():
            content2 = row[1]
            if imghdr.what(None, h=content2) is not None:
                continue
            if content2.startswith(b'%PDF'):
                pdf = pdfplumber.open(io.BytesIO(content2))
                extracted_text = ""
                for page in pdf.pages:
                    extracted_text += page.extract_text()
                paragraf_docx2 = extracted_text.lower()
                print(paragraf_docx2)
            else:
                paragraf_docx2 = ""
                document = Document(io.BytesIO(content2))
                for paragraph in document.paragraphs:
                    paragraf_docx2 += paragraph.text.lower()
            sentences1 = preprocess_sentences(paragraf_docx)
            sentences2 = preprocess_sentences(paragraf_docx2)
            arrteks=[]
            arrteks2=[]
            arrResHash = set()
            for sentence in sentences1:
                arrteks.append(sentence)
                sentence_hash = hashlib.sha256(sentence.encode()).hexdigest()
                arrResHash.add(sentence_hash)
            
            arrResHash2 = set()
            for sentence in sentences2:
                arrteks2.append(sentence)
                sentence_hash = hashlib.sha256(sentence.encode()).hexdigest()
                arrResHash2.add(sentence_hash)

            result_jaccard = jaccard_similarity(arrResHash, arrResHash2) * 100
            jaccard_dec = np.around(result_jaccard, decimals=4)
            final_result = '{:.4f}%'.format(jaccard_dec)

            result_average = average_similarity(arrResHash, arrResHash2) * 100
            average_dec = np.around(result_average, decimals=4)
            final_result_average = '{:.4f}%'.format(average_dec)
            
            similarity_results_jaccard.append(f"Jaccard similarity with document id {row[0]}: {final_result}")

            similarity_results_average.append(f"Average similarity with document id {row[0]}: {final_result}")
            # print("arr1")
            # print(sentences1)
            # print("arr2")
            # print (sentences2)
            if result_jaccard > 0 and result_average > 0:
        
                for i, similarity in enumerate(arrResHash):
                    for j, simi2 in enumerate(arrResHash2):
                        if similarity == simi2:
                            similarity_results_average.append(f"Kemiripan pada Baris {j+1} Isi: {arrteks2[j]}")

        # return "<br>".join(lines_with_similarity)

        return "<br><br>".join(similarity_results_jaccard).join(similarity_results_average)
    else:
        return 'Invalid document ID'
    

@app.route('/compare_images', methods=['GET','POST'])
def calculate_images():
    doc_id = session['doc_id']
    # Retrieve the image blobs from the database
    cursor = mysql.connection.cursor()
    cursor.execute("SELECT image_file, image_name FROM image_doc WHERE img_id != %s", (doc_id,))
    rows = cursor.fetchall()

    results=[]
    # Load the image to compare
    # Load the image to compare
   
    cursor.execute("SELECT image_file FROM image_doc WHERE img_id = %s", (doc_id,))
    img_to_compare_blob = cursor.fetchone()[0]

    # Convert the blob data to a NumPy array
    nparr = np.frombuffer(img_to_compare_blob, np.uint8)

    # Decode the NumPy array as an image
    img_to_compare = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

    # Convert the image to grayscale
    gray_img_to_compare = cv2.cvtColor(img_to_compare, cv2.COLOR_BGR2GRAY)

    # Calculate the histogram of the image to compare
    hist_to_compare = cv2.calcHist([gray_img_to_compare], [0], None, [256], [0, 256])
    cv2.normalize(hist_to_compare, hist_to_compare, alpha=0, beta=1, norm_type=cv2.NORM_MINMAX)

    total_score = 0
    total_images = 0

    # Loop through each row in the database
    for row in rows:
        blob_data = row[0]

        if imghdr.what(None, h=blob_data) is not None:
            # Convert the blob data to a NumPy array
            nparr = np.frombuffer(blob_data, np.uint8)

            # Decode the NumPy array as an image
            dataset_img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

            # Convert the dataset image to grayscale
            gray_dataset_img = cv2.cvtColor(dataset_img, cv2.COLOR_BGR2GRAY)

            # Calculate the histogram of the dataset image
            hist_dataset_img = cv2.calcHist([gray_dataset_img], [0], None, [256], [0, 256])
            cv2.normalize(hist_dataset_img, hist_dataset_img, alpha=0, beta=1, norm_type=cv2.NORM_MINMAX)

            # Compare the histograms of the two images
            score = cv2.compareHist(hist_to_compare, hist_dataset_img, cv2.HISTCMP_CORREL)

            # Convert the score to a percentage
            percentage_score = score * 100
            total_score += percentage_score
            total_images += 1

            # Append the percentage score to the results list
            results.append(f"Image Similarity with Image name {row[1]}: {percentage_score:.2f}")
    
        # Calculate the average plagiarism score
        average_score = total_score / total_images

        # Append the average score to the results list
        results.append(f"Average Plagiarism Score: {average_score:.2f}")

        # # Return the results as a response
        # return "<br>".join(results)

    # Return the results as a response
    return "<br>".join(results)



if __name__ == '__main__':
    app.run(host='localhost', port=5000, debug=True)


    
    